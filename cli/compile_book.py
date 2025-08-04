import argparse
import os
import re
import json
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from config.settings import settings
from infrastructure.llm_client import GeminiClient
from infrastructure.outline_manager import OutlineManager

def generate_book_title(summaries: str, language=settings.DEFAULT_LANGUAGE) -> list:
    """Генерирует варианты названия книги на основе саммари глав"""
    llm = GeminiClient(language)
    
    prompt = f"""
    [Language: {language}]
    Based on these chapter summaries of a book, generate 10 creative and catchy book titles.
    The titles should be relevant to the overall story and themes.
    
    Chapter Summaries:
    {summaries}
    
    Return ONLY a JSON array with 10 title strings:
    ["Title 1", "Title 2", ...]
    """
    
    result = llm.generate_text(prompt)
    
    try:
        # Удаляем возможные обертки кода
        clean_result = re.sub(r'```(json)?\n|\n```', '', result).strip()
        titles = json.loads(clean_result)
        return titles[:10]  # Берем первые 10 вариантов
    except (json.JSONDecodeError, TypeError) as e:
        print(f"Failed to parse titles: {e}")
        print(f"Response: {result[:500]}...")
        return ["My Generated Book"] * 10  # Запасной вариант

def compile_book(language: str, custom_title: str = None):
    """Собирает книгу из сгенерированных глав в DOCX"""
    # Создаем директорию для книг
    Path(settings.BOOKS_DIR).mkdir(exist_ok=True)
    
    manager = OutlineManager()
    df = manager.load_outline()
    
    # Собираем все саммари для генерации названия
    summaries = "\n".join(
        f"Chapter {row['Chapter']}: {row['Summary']}" 
        for _, row in df.iterrows() 
        if not pd.isna(row["Summary"])
    )
    
    # Генерируем варианты названия
    if not custom_title:
        titles = generate_book_title(summaries)
        print("\nGenerated book titles:")
        for i, title in enumerate(titles, 1):
            print(f"{i}. {title}")
        
        # Выбираем название по умолчанию (первое)
        book_title = titles[0]
        print(f"\nUsing default title: '{book_title}'")
    else:
        book_title = custom_title
        print(f"Using custom title: '{book_title}'")
    
    # Создаем документ
    doc = Document()
    
    # Настройка стилей
    styles = doc.styles
    title_style = styles["Title"]
    title_font = title_style.font
    title_font.name = "Times New Roman"
    title_font.size = Pt(24)
    
    heading_style = styles["Heading 1"]
    heading_font = heading_style.font
    heading_font.name = "Times New Roman"
    heading_font.size = Pt(16)
    heading_font.bold = True
    
    normal_style = styles["Normal"]
    normal_font = normal_style.font
    normal_font.name = "Times New Roman"
    normal_font.size = Pt(12)
    
    # Добавляем заголовок книги
    title_para = doc.add_paragraph(book_title, style="Title")
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()
    
    # Сортируем главы по номеру
    df = df.sort_values(by="Chapter")
    
    # Добавляем главы
    for _, row in df.iterrows():
        if pd.isna(row["File"]) or not row["File"]:
            continue
            
        chapter_file = row["File"]
        if not os.path.exists(chapter_file):
            print(f"Warning: Chapter file not found: {chapter_file}")
            continue
            
        # Добавляем заголовок главы
        # chapter_title = f"Chapter {int(row['Chapter'])}: {row['Title']}"
        chapter_title = f"{row['Title']}"
        doc.add_paragraph(chapter_title, style="Heading 1")
        
        # Читаем и добавляем текст главы
        with open(chapter_file, "r", encoding="utf-8") as f:
            chapter_text = f.read()
        
        # Разделяем текст на параграфы
        paragraphs = chapter_text.split("\n\n")
        for para in paragraphs:
            if para.strip():  # Пропускаем пустые параграфы
                doc.add_paragraph(para.strip(), style="Normal")
        
        doc.add_page_break()
    
    # Сохраняем книгу
    safe_title = re.sub(r'[^\w\s-]', '', book_title)[:50]
    filename = f"{settings.BOOKS_DIR}/{safe_title}.docx"
    doc.save(filename)
    print(f"\nBook compiled successfully: {filename}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Compile book from generated chapters")
    parser.add_argument("--language", default=settings.DEFAULT_LANGUAGE, help="Book language")
    parser.add_argument("--title", help="Custom book title (optional)")
    args = parser.parse_args()
    
    compile_book(args.language, args.title)